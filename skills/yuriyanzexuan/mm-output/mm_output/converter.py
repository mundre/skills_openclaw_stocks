"""
Core converter module for multi-modal output generation.
"""

import os
import re
import io
import base64
from pathlib import Path
from typing import Optional, Union, List
from urllib.parse import urljoin, urlparse


class MMOutputGenerator:
    """
    Multi-modal output generator for converting HTML to PDF, PNG, and DOCX formats.
    
    Uses headless Chrome (via Playwright) for PDF/PNG generation,
    and python-docx for DOCX generation.
    """
    
    def __init__(self, chrome_path: Optional[str] = None):
        """
        Initialize the MMOutputGenerator.
        
        Args:
            chrome_path: Optional path to Chrome/Chromium executable.
                        If not provided, will try to auto-detect or use playwright's chromium.
                        Can also be set via CHROME_EXECUTABLE_PATH environment variable.
        """
        import os
        self.name = "MMOutputGenerator"
        # Priority: constructor arg > env var > auto-detect
        self._chrome_path = chrome_path or os.getenv("CHROME_EXECUTABLE_PATH")
        self._playwright = None
        self._browser = None
        
    def _ensure_playwright(self):
        """Ensure playwright is installed and browser is launched."""
        if self._browser is not None:
            return
            
        try:
            from playwright.sync_api import sync_playwright
        except ImportError:
            raise ImportError(
                "playwright is required for PDF/PNG generation. "
                "Install with: pip install playwright && playwright install chromium"
            )
        
        self._playwright = sync_playwright().start()
        
        launch_options = {"headless": True}
        if self._chrome_path:
            launch_options["executable_path"] = self._chrome_path
            print(f"[{self.name}] Using custom Chrome: {self._chrome_path}")
        else:
            # Try to auto-detect Chrome/Chromium
            chrome_paths = [
                # Linux
                "/usr/bin/google-chrome",
                "/usr/bin/chromium",
                "/usr/bin/chromium-browser",
                # Custom location
                "/mnt/tidalfs-bdsz01/usr/tusen/yanzexuan/postergenparserunit/chrome-linux64/chrome",
            ]
            for path in chrome_paths:
                if Path(path).exists():
                    launch_options["executable_path"] = path
                    print(f"[{self.name}] Auto-detected Chrome: {path}")
                    break
            
        self._browser = self._playwright.chromium.launch(**launch_options)
        
    def _close_browser(self):
        """Close browser and cleanup playwright."""
        if self._browser:
            self._browser.close()
            self._browser = None
        if self._playwright:
            self._playwright.stop()
            self._playwright = None
            
    def __enter__(self):
        """Context manager entry."""
        return self
        
    def __exit__(self, exc_type, exc_val, exc_tb):
        """Context manager exit."""
        self._close_browser()
        
    def html_to_pdf(
        self,
        html_path: Union[str, Path],
        output_path: Union[str, Path],
        page_size: str = "A4",
        margin: dict = None,
        landscape: bool = False,
        print_background: bool = True
    ) -> str:
        """
        Convert HTML file to PDF using headless Chrome.
        
        Args:
            html_path: Path to input HTML file
            output_path: Path to output PDF file
            page_size: Page size (A4, Letter, Legal, etc.)
            margin: Page margins in dict format {"top": "1cm", "right": "1cm", ...}
            landscape: Whether to use landscape orientation
            print_background: Whether to print background graphics
            
        Returns:
            Path to generated PDF file
        """
        self._ensure_playwright()
        
        html_path = Path(html_path).resolve()
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Create context with proper viewport
        context_options = {"viewport": {"width": 1200, "height": 1600}}
        context = self._browser.new_context(**context_options)
        
        try:
            page = context.new_page()
            
            # Load HTML file
            page.goto(f"file://{html_path}", wait_until="networkidle")
            
            # Wait for fonts and images to load
            page.wait_for_timeout(2000)
            
            # Prepare margin settings
            default_margin = {"top": "1cm", "right": "1cm", "bottom": "1cm", "left": "1cm"}
            if margin:
                default_margin.update(margin)
            
            # Generate PDF
            pdf_options = {
                "path": str(output_path),
                "format": page_size,
                "margin": default_margin,
                "print_background": print_background,
                "landscape": landscape,
            }
            
            page.pdf(**pdf_options)
            
            print(f"[{self.name}] PDF generated: {output_path}")
            return str(output_path)
            
        finally:
            context.close()
            
    def html_to_png(
        self,
        html_path: Union[str, Path],
        output_path: Union[str, Path],
        full_page: bool = True,
        viewport_size: tuple = (1200, 1600),
        wait_time: int = 2000
    ) -> str:
        """
        Convert HTML file to PNG image using headless Chrome screenshot.
        
        Args:
            html_path: Path to input HTML file
            output_path: Path to output PNG file
            full_page: Whether to capture full page or just viewport
            viewport_size: Viewport size as (width, height) tuple
            wait_time: Time to wait for page load in milliseconds
            
        Returns:
            Path to generated PNG file
        """
        self._ensure_playwright()
        
        html_path = Path(html_path).resolve()
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        context = self._browser.new_context(
            viewport={"width": viewport_size[0], "height": viewport_size[1]}
        )
        
        try:
            page = context.new_page()
            
            # Load HTML file
            page.goto(f"file://{html_path}", wait_until="networkidle")
            
            # Wait for fonts and images
            page.wait_for_timeout(wait_time)
            
            # Take screenshot
            screenshot_options = {
                "path": str(output_path),
                "full_page": full_page,
                "type": "png"
            }
            
            page.screenshot(**screenshot_options)
            
            print(f"[{self.name}] PNG generated: {output_path}")
            return str(output_path)
            
        finally:
            context.close()
            
    def html_to_docx(
        self,
        html_path: Union[str, Path],
        output_path: Union[str, Path],
        include_images: bool = True
    ) -> str:
        """
        Convert HTML file to DOCX using python-docx.
        
        Note: This is a best-effort conversion that extracts text content
        and basic formatting. Complex layouts may not be perfectly preserved.
        
        Args:
            html_path: Path to input HTML file
            output_path: Path to output DOCX file
            include_images: Whether to include images in the document
            
        Returns:
            Path to generated DOCX file
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX generation. "
                "Install with: pip install python-docx"
            )
        
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError(
                "beautifulsoup4 is required for HTML parsing. "
                "Install with: pip install beautifulsoup4"
            )
        
        html_path = Path(html_path).resolve()
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Read and parse HTML
        html_content = html_path.read_text(encoding='utf-8')
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Create document
        doc = Document()
        
        # Helper function to add styled text
        def add_styled_text(paragraph, text, bold=False, italic=False, size=None, color=None):
            run = paragraph.add_run(text)
            run.bold = bold
            run.italic = italic
            if size:
                run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor(*color)
            return run
        
        # Extract and convert content
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
            
        # Process body content
        body = soup.find('body') or soup
        
        for element in body.descendants:
            if element.name is None:
                continue
                
            # Handle headings
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1])
                p = doc.add_heading(level=min(level, 9))
                p.add_run(element.get_text(strip=True))
                
            # Handle paragraphs
            elif element.name == 'p':
                text = element.get_text(strip=True)
                if text:
                    doc.add_paragraph(text)
                    
            # Handle images
            elif element.name == 'img' and include_images:
                img_src = element.get('src', '')
                if img_src:
                    # Resolve relative paths
                    if not img_src.startswith(('http://', 'https://', 'data:')):
                        img_path = html_path.parent / img_src
                        if img_path.exists():
                            try:
                                doc.add_picture(str(img_path), width=Inches(5.0))
                            except Exception as e:
                                print(f"[{self.name}] Warning: Could not add image {img_src}: {e}")
                                
            # Handle lists
            elif element.name == 'li':
                # Check if it's an ordered or unordered list
                parent = element.find_parent(['ul', 'ol'])
                if parent:
                    text = element.get_text(strip=True)
                    if text:
                        style = 'List Number' if parent.name == 'ol' else 'List Bullet'
                        doc.add_paragraph(text, style=style)
                        
            # Handle tables (basic support)
            elif element.name == 'table':
                rows = element.find_all('tr')
                if rows:
                    # Count columns from first row
                    first_row = rows[0]
                    cols = len(first_row.find_all(['td', 'th']))
                    if cols > 0:
                        table = doc.add_table(rows=len(rows), cols=cols)
                        table.style = 'Light Grid Accent 1'
                        
                        for i, row in enumerate(rows):
                            cells = row.find_all(['td', 'th'])
                            for j, cell in enumerate(cells):
                                if j < cols:
                                    table.rows[i].cells[j].text = cell.get_text(strip=True)
                                    
        # Save document
        doc.save(str(output_path))
        
        print(f"[{self.name}] DOCX generated: {output_path}")
        return str(output_path)
        
    def convert_all(
        self,
        html_path: Union[str, Path],
        output_dir: Union[str, Path],
        base_name: Optional[str] = None,
        formats: Optional[List[str]] = None
    ) -> dict:
        """
        Convert HTML to all supported formats.
        
        Args:
            html_path: Path to input HTML file
            output_dir: Directory to save output files
            base_name: Base filename without extension (defaults to input filename)
            formats: List of formats to generate ['pdf', 'png', 'docx'] or None for all
            
        Returns:
            Dictionary mapping format names to output file paths
        """
        html_path = Path(html_path)
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        if base_name is None:
            base_name = html_path.stem
            
        formats = formats or ['pdf', 'png', 'docx']
        results = {}
        
        try:
            for fmt in formats:
                fmt = fmt.lower()
                output_path = output_dir / f"{base_name}.{fmt}"
                
                if fmt == 'pdf':
                    results['pdf'] = self.html_to_pdf(html_path, output_path)
                elif fmt == 'png':
                    results['png'] = self.html_to_png(html_path, output_path)
                elif fmt == 'docx':
                    results['docx'] = self.html_to_docx(html_path, output_path)
                else:
                    print(f"[{self.name}] Warning: Unsupported format '{fmt}', skipping")
                    
        finally:
            # Cleanup browser resources
            self._close_browser()
            
        return results


# Convenience functions for direct use
def html_to_pdf(html_path: Union[str, Path], output_path: Union[str, Path], **kwargs) -> str:
    """Standalone function to convert HTML to PDF."""
    with MMOutputGenerator() as gen:
        return gen.html_to_pdf(html_path, output_path, **kwargs)


def html_to_png(html_path: Union[str, Path], output_path: Union[str, Path], **kwargs) -> str:
    """Standalone function to convert HTML to PNG."""
    with MMOutputGenerator() as gen:
        return gen.html_to_png(html_path, output_path, **kwargs)


def html_to_docx(html_path: Union[str, Path], output_path: Union[str, Path], **kwargs) -> str:
    """Standalone function to convert HTML to DOCX."""
    with MMOutputGenerator() as gen:
        return gen.html_to_docx(html_path, output_path, **kwargs)


def convert_all(html_path: Union[str, Path], output_dir: Union[str, Path], **kwargs) -> dict:
    """Standalone function to convert HTML to all formats."""
    with MMOutputGenerator() as gen:
        return gen.convert_all(html_path, output_dir, **kwargs)
