#!/usr/bin/env python3
"""Index local files into ChromaDB with BGE-M3 embeddings and parent-child chunking.

Parent chunks (large) are stored as documents for retrieval.
Child chunks (small) are embedded for precise semantic search.
When a child matches, the parent is returned for full context.

Usage:
    index.py [--reindex] [--paths PATH1 PATH2 ...]

Default paths: ~/Documenti and ~/Scaricati.
Use --reindex to drop existing data and re-index from scratch.
"""

import argparse
import hashlib
import logging
import os
import sys
from pathlib import Path

import chromadb
from sentence_transformers import SentenceTransformer

# Silence pdfminer FontBBox spam
logging.getLogger("pdfminer").setLevel(logging.ERROR)

# --- Config ---

EMBED_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
PARENT_SIZE = 768       # words per parent chunk
CHILD_SIZE = 128        # words per child chunk
CHILD_OVERLAP = 24      # overlap between children
MAX_FILE_SIZE_MB = 30
BATCH_SIZE = 1          # embedding batch size (conservative for 4GB RAM)

# Security: only allow indexing under these directories
ALLOWED_ROOTS = [
    os.path.expanduser("~/Documenti/github/thesis"),
    os.path.expanduser("~/Documenti/github/polito"),
    os.path.expanduser("~/Documenti"),
    os.path.expanduser("~/Scaricati"),
]

BLOCKED_PATTERNS = {".ssh", ".gnupg", ".env", "credentials", "tokens", ".config/openclaw"}

DEFAULT_PATHS = [
    "~/Documenti",
    "~/Scaricati",
]

TEXT_EXTENSIONS = {
    ".txt", ".md", ".rst", ".csv", ".tsv",
    ".yaml", ".yml", ".json", ".toml", ".cfg", ".ini", ".xml", ".html", ".css",
    ".tex", ".bib", ".log",
}

DOCUMENT_EXTENSIONS = {
    ".pdf",
    ".docx", ".odt", ".rtf", ".epub", ".xlsx", ".pptx",
}

ALL_EXTENSIONS = TEXT_EXTENSIONS | DOCUMENT_EXTENSIONS

DB_DIR = os.path.expanduser("~/.local/share/local-rag/chromadb")


# --- Helpers ---

def file_hash(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def split_words(text: str) -> list[str]:
    return text.split()


def make_parent_chunks(words: list[str], size: int = PARENT_SIZE) -> list[str]:
    """Split words into parent chunks (no overlap)."""
    parents = []
    for i in range(0, len(words), size):
        chunk = " ".join(words[i : i + size])
        if chunk.strip():
            parents.append(chunk)
    return parents


def make_child_chunks(parent_words: list[str], size: int = CHILD_SIZE, overlap: int = CHILD_OVERLAP) -> list[str]:
    """Split a parent's words into overlapping child chunks."""
    children = []
    for i in range(0, len(parent_words), size - overlap):
        chunk = " ".join(parent_words[i : i + size])
        if chunk.strip():
            children.append(chunk)
    return children


def extract_text(filepath: str) -> str | None:
    ext = Path(filepath).suffix.lower()

    if ext in TEXT_EXTENSIONS:
        try:
            return Path(filepath).read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return None

    if ext == ".pdf":
        try:
            # Skip PDFs > 5MB to avoid OOM on 4GB RAM
            if os.path.getsize(filepath) > 5 * 1024 * 1024:
                print(f"    Skipping {filepath}: PDF too large for RAM ({os.path.getsize(filepath)//1024//1024}MB > 5MB limit)")
                return None
            from pdfminer.high_level import extract_text
            text = extract_text(filepath) or ""
            if len(text) > 5 * 1024 * 1024:
                print(f"    Skipping {filepath}: extracted text too large ({len(text)//1024//1024}MB)")
                return None
            return text
        except Exception:
            return None

    if ext == ".docx":
        try:
            from docx import Document
            return "\n".join(p.text for p in Document(filepath).paragraphs)
        except ImportError:
            import subprocess
            subprocess.check_call([
                "uv", "pip", "install", "-p", sys.executable, "python-docx",
            ])
            from docx import Document
            return "\n".join(p.text for p in Document(filepath).paragraphs)
        except Exception:
            return None

    return None


def crawl(root: str) -> list[str]:
    files = []
    root = os.path.expanduser(root)
    if not os.path.isdir(root):
        print(f"  Skipping {root} (not a directory)")
        return files
    # Security: validate path is under allowed roots
    root_real = os.path.realpath(root)
    allowed = any(root_real.startswith(os.path.realpath(ar)) for ar in ALLOWED_ROOTS)
    if not allowed:
        print(f"  Skipping {root} (outside allowed directories)")
        return files
    for dirpath, dirnames, filenames in os.walk(root):
        # Security: skip sensitive directories
        dirnames[:] = [d for d in dirnames
                       if not d.startswith(".")
                       and d not in {
                           "node_modules", "__pycache__", ".git", ".venv", "venv",
                           "labs", "exercises", *BLOCKED_PATTERNS,
                       }]
        for fn in filenames:
            ext = Path(fn).suffix.lower()
            if ext in ALL_EXTENSIONS:
                full = os.path.join(dirpath, fn)
                size_mb = os.path.getsize(full) / (1024 * 1024)
                if size_mb <= MAX_FILE_SIZE_MB:
                    files.append(full)
    return files


# --- Main ---

LOCK_FILE = os.path.expanduser("~/.local/share/local-rag/index.lock")


def main():
    # Prevent concurrent runs
    import fcntl
    os.makedirs(os.path.dirname(LOCK_FILE), exist_ok=True)
    lock_fd = open(LOCK_FILE, "w")
    try:
        fcntl.flock(lock_fd, fcntl.LOCK_EX | fcntl.LOCK_NB)
    except OSError:
        print("Another indexing process is already running. Exiting.")
        sys.exit(0)

    parser = argparse.ArgumentParser(description="Index local files with parent-child chunking")
    parser.add_argument("--reindex", action="store_true", help="Drop existing data and re-index")
    parser.add_argument("--paths", nargs="+", default=DEFAULT_PATHS, help="Directories to index")
    args = parser.parse_args()

    print(f"Loading embedding model: {EMBED_MODEL}...")
    model = SentenceTransformer(EMBED_MODEL, trust_remote_code=True)

    print(f"Opening ChromaDB at {DB_DIR}...")
    client = chromadb.PersistentClient(path=DB_DIR)

    if args.reindex:
        print("Dropping existing collections...")
        for name in ["parents", "children"]:
            try:
                client.delete_collection(name)
            except Exception:
                pass

    parents_col = client.get_or_create_collection("parents", metadata={"hnsw:space": "cosine"})
    children_col = client.get_or_create_collection("children", metadata={"hnsw:space": "cosine"})

    total_parents = 0
    total_children = 0
    total_files = 0

    for root in args.paths:
        print(f"\nScanning {root}...")
        files = crawl(root)
        print(f"  Found {len(files)} files to index")

        for filepath in files:
            text = extract_text(filepath)
            if not text or not text.strip():
                continue

            fhash = file_hash(filepath)

            # Check if already indexed with same hash
            existing = parents_col.get(
                where={"filepath": filepath},
                include=["metadatas"],
                limit=1,
            )
            if existing["metadatas"]:
                old_hash = existing["metadatas"][0].get("file_hash")
                if old_hash == fhash:
                    continue  # Unchanged

                # File changed — remove old chunks
                parents_col.delete(where={"filepath": filepath})
                children_col.delete(where={"filepath": filepath})

            words = split_words(text)
            parents = make_parent_chunks(words)

            all_child_texts = []
            all_child_metas = []
            all_child_ids = []
            parent_ids = []
            parent_docs = []
            parent_metas = []

            for p_idx, parent_text in enumerate(parents):
                parent_id = f"{fhash}_p{p_idx}"
                parent_ids.append(parent_id)
                parent_docs.append(parent_text)
                parent_metas.append({
                    "filepath": filepath,
                    "parent_idx": p_idx,
                    "file_hash": fhash,
                })

                p_words = split_words(parent_text)
                child_texts = make_child_chunks(p_words)
                for c_idx, child_text in enumerate(child_texts):
                    child_id = f"{fhash}_p{p_idx}_c{c_idx}"
                    all_child_ids.append(child_id)
                    all_child_texts.append(child_text)
                    all_child_metas.append({
                        "filepath": filepath,
                        "parent_idx": p_idx,
                        "child_idx": c_idx,
                        "file_hash": fhash,
                    })

            # Store parents (no embeddings needed — just text storage)
            parents_col.add(ids=parent_ids, documents=parent_docs, metadatas=parent_metas)

            # Embed and store children in batches
            for batch_start in range(0, len(all_child_texts), BATCH_SIZE):
                batch_texts = all_child_texts[batch_start : batch_start + BATCH_SIZE]
                batch_ids = all_child_ids[batch_start : batch_start + BATCH_SIZE]
                batch_metas = all_child_metas[batch_start : batch_start + BATCH_SIZE]

                embeddings = model.encode(batch_texts, show_progress_bar=False).tolist()
                children_col.add(ids=batch_ids, embeddings=embeddings, documents=batch_texts, metadatas=batch_metas)

            total_parents += len(parent_ids)
            total_children += len(all_child_ids)
            total_files += 1

            if total_files % 50 == 0:
                print(f"  Progress: {total_files} files, {total_parents} parents, {total_children} children")

    print(f"\nDone! Indexed {total_files} files.")
    print(f"  Parents: {total_parents} (stored as full context)")
    print(f"  Children: {total_children} (embedded for search)")
    print(f"  Parents collection: {parents_col.count()} chunks")
    print(f"  Children collection: {children_col.count()} chunks")


if __name__ == "__main__":
    main()
