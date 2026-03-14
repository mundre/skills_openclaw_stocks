#!/usr/bin/env python
"""Build a technical accounting report DOCX from JSON input."""

from __future__ import annotations

import argparse
import json
from datetime import date
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.shared import Pt
except ImportError as exc:  # pragma: no cover
    raise SystemExit(
        "python-docx is required. Install with: python -m pip install --user python-docx"
    ) from exc


def _as_text(value: Any, default: str = "") -> str:
    if value is None:
        return default
    if isinstance(value, str):
        return value.strip()
    return str(value).strip()


def _as_list_of_str(value: Any) -> list[str]:
    if isinstance(value, list):
        return [_as_text(item) for item in value if _as_text(item)]
    if isinstance(value, str) and value.strip():
        return [value.strip()]
    return []


def _as_list_of_dict(value: Any) -> list[dict[str, Any]]:
    if not isinstance(value, list):
        return []
    out: list[dict[str, Any]] = []
    for item in value:
        if isinstance(item, dict):
            out.append(item)
    return out


def _set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _add_metadata_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for index, (label, value) in enumerate(rows):
        table.cell(index, 0).text = label
        table.cell(index, 1).text = value


def _add_heading_and_text(doc: Document, heading: str, text: str) -> None:
    if not text:
        return
    doc.add_heading(heading, level=1)
    doc.add_paragraph(text)


def _add_heading_and_bullets(doc: Document, heading: str, items: list[str]) -> None:
    if not items:
        return
    doc.add_heading(heading, level=1)
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _add_guidance_table(doc: Document, guidance: list[dict[str, Any]]) -> None:
    if not guidance:
        return

    doc.add_heading("Guidance Reviewed", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Citation"
    header[1].text = "Type"
    header[2].text = "Key Point"
    header[3].text = "URL"
    header[4].text = "Accessed"

    for item in guidance:
        row = table.add_row().cells
        row[0].text = _as_text(item.get("citation"))
        row[1].text = _as_text(item.get("source_type"))
        row[2].text = _as_text(item.get("key_point"))
        row[3].text = _as_text(item.get("url"))
        row[4].text = _as_text(item.get("accessed_on"))


def _add_journal_entries_table(doc: Document, entries: list[dict[str, Any]]) -> None:
    if not entries:
        return

    doc.add_heading("Journal Entry Examples", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Description"
    header[1].text = "Debit"
    header[2].text = "Credit"
    header[3].text = "Amount"

    for item in entries:
        row = table.add_row().cells
        row[0].text = _as_text(item.get("description"))
        row[1].text = _as_text(item.get("debit"))
        row[2].text = _as_text(item.get("credit"))
        row[3].text = _as_text(item.get("amount"))


def _add_qa(doc: Document, qa_items: list[dict[str, Any]]) -> None:
    if not qa_items:
        return

    doc.add_heading("Q and A", level=1)
    for item in qa_items:
        question = _as_text(item.get("question"))
        answer = _as_text(item.get("answer"))
        if question:
            p = doc.add_paragraph()
            p.add_run("Q: ").bold = True
            p.add_run(question)
        if answer:
            p = doc.add_paragraph()
            p.add_run("A: ").bold = True
            p.add_run(answer)


def _build_document(payload: dict[str, Any], report_format: str) -> Document:
    doc = Document()
    _set_default_style(doc)

    default_title_map = {
        "memo": "Technical Accounting Memorandum",
        "email": "Technical Accounting Email Draft",
        "q-and-a": "Technical Accounting Q and A",
    }
    title = _as_text(payload.get("title"), default_title_map[report_format])
    doc.add_heading(title, level=0)

    report_date = _as_text(payload.get("date"), date.today().isoformat())

    if report_format == "memo":
        _add_metadata_table(
            doc,
            [
                ("To", _as_text(payload.get("prepared_for"))),
                ("From", _as_text(payload.get("prepared_by"))),
                ("Date", report_date),
                ("Subject", _as_text(payload.get("subject"))),
            ],
        )
    elif report_format == "email":
        _add_metadata_table(
            doc,
            [
                ("To", _as_text(payload.get("to"), _as_text(payload.get("prepared_for")))),
                ("Cc", _as_text(payload.get("cc"))),
                ("From", _as_text(payload.get("from"), _as_text(payload.get("prepared_by")))),
                ("Date", report_date),
                ("Subject", _as_text(payload.get("subject"))),
            ],
        )
    else:
        _add_metadata_table(
            doc,
            [
                ("Prepared For", _as_text(payload.get("prepared_for"))),
                ("Prepared By", _as_text(payload.get("prepared_by"))),
                ("Date", report_date),
                ("Topic", _as_text(payload.get("subject"))),
            ],
        )

    _add_heading_and_text(doc, "Issue", _as_text(payload.get("issue")))
    _add_heading_and_bullets(doc, "Facts", _as_list_of_str(payload.get("facts")))

    guidance = _as_list_of_dict(payload.get("guidance"))
    _add_guidance_table(doc, guidance)

    _add_heading_and_bullets(doc, "Analysis", _as_list_of_str(payload.get("analysis")))
    _add_heading_and_text(doc, "Conclusion", _as_text(payload.get("conclusion")))

    _add_heading_and_bullets(
        doc,
        "Disclosure Considerations",
        _as_list_of_str(payload.get("disclosure_considerations")),
    )

    _add_journal_entries_table(doc, _as_list_of_dict(payload.get("journal_entries")))

    _add_heading_and_bullets(doc, "Assumptions", _as_list_of_str(payload.get("assumptions")))
    _add_heading_and_bullets(doc, "Open Items", _as_list_of_str(payload.get("open_items")))
    _add_heading_and_bullets(doc, "Next Steps", _as_list_of_str(payload.get("next_steps")))

    _add_qa(doc, _as_list_of_dict(payload.get("qa")))

    if guidance:
        doc.add_heading("Source List", level=1)
        for source in guidance:
            citation = _as_text(source.get("citation"), "Unlabeled source")
            source_type = _as_text(source.get("source_type"), "unknown type")
            accessed_on = _as_text(source.get("accessed_on"), "unknown date")
            url = _as_text(source.get("url"), "(no URL)")
            doc.add_paragraph(
                f"{citation} ({source_type}), accessed {accessed_on}: {url}",
                style="List Bullet",
            )

    return doc


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate a technical accounting DOCX report from JSON input."
    )
    parser.add_argument("--input-json", required=True, help="Path to JSON input payload")
    parser.add_argument("--output-docx", required=True, help="Path for generated DOCX")
    parser.add_argument(
        "--format",
        default="memo",
        choices=["memo", "email", "q-and-a"],
        help="Output report style",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    input_path = Path(args.input_json)
    output_path = Path(args.output_docx)

    with input_path.open("r", encoding="utf-8") as f:
        payload = json.load(f)

    if not isinstance(payload, dict):
        raise SystemExit("Input JSON must contain an object at top level.")

    doc = _build_document(payload, args.format)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    print(f"[OK] Generated {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
